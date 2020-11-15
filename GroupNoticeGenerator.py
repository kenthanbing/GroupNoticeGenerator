import docx,os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.oxml.ns import qn
import json

doc = Document('GroupNoticeGenerator' + os.sep + 'letterhead.docx')

#设置run字体
def set_font(run):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return

# 定义写入段落的函数
def write_paragraph(paragraphs):
    for paragraph in paragraphs:
        paragraph = doc.add_paragraph(paragraph)
        paragraph.paragraph_format.first_line_indent = Inches(0.32)
    return    

# 获取json文件数据
with open('GroupNoticeGenerator' + os.sep + 'notice_data.json', 'rb') as f:
    notice_data = json.load(f)

# 设置正文全局字体
doc.styles['Normal'].font.name = '微软雅黑'
doc.styles['Normal'].font.size = Pt(12)
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 写标题
doc.add_heading(level=0).add_run(notice_data['exhibition_title'])
doc.add_heading(level=0).add_run('展团参展指南')

# 写参展人员安排
doc.add_heading(level=3).add_run('一、参展人员安排')
doc.add_paragraph('1、酒店名称：')
doc.add_paragraph(f'   {notice_data["hotel"]}')
doc.add_paragraph('2、展团领队：')
doc.add_paragraph(f'   {notice_data["group_leader"]}')
doc.add_paragraph('3、集合时间及地点：')
doc.add_paragraph(f'   {notice_data["time"]} {notice_data["place"]}处集合。领队将举导游旗在集合处等候。')

# 写观展及旅游须知
doc.add_heading(level=3).add_run('二、观展及旅游须知')
doc.add_paragraph('   一、展览会及旅游期间的详细情况详见《日程安排》。展团需同进同出，请大家按时参加各项活动，有事或者需要拜访客户需要告知领队，并在离团申请书上签字。')
doc.add_paragraph('   二、观展期间由领队带领进出展馆，遵守展馆纪律，尊重他国风俗，保持良好的国际形象。')
doc.add_paragraph('   三、注意安全，妥善保管好参展样品及个人携带物品。建议将重要物品存放在酒店房间保险柜，不要随身携带，也不要带入会场。贵重物品不要随手乱放，也不要交给陌生人看管。')
doc.add_paragraph('   四、用餐、坐车及活动内容请遵照《日程安排》。在酒店需要观看收费电视节目和打电话，请自行到总台开通，并结清。个人在酒店的杂费消费需在退房前提前结清，住房名单如需调换请与领队联系。')
doc.add_paragraph('   其他特别须知：')

paragraphs = notice_data['must_know']
write_paragraph(paragraphs)

# 写展会信息
doc.add_heading(level=3).add_run('三、展会信息')
doc.add_paragraph('1、展会日程安排：')
doc.add_paragraph(f'   {notice_data["prepare_time"]}')
doc.add_paragraph(f'   {notice_data["open_time"]}')
doc.add_paragraph('2、展馆名称及地址：')
doc.add_paragraph(f'   {notice_data["venue_name"]}')
doc.add_paragraph(f'   地址：{notice_data["venue_address"]}')

# 写附件
doc.add_paragraph()
doc.add_paragraph('附：行程')


doc.save('GroupNoticeGenerator' + os.sep + f'{notice_data["exhibition_title"]}展团参展指南.docx')